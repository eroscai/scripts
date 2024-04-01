from docx import Document
from docx.shared import Pt
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import random

def generate_two_digit_plus_minus_one_digit(total_problems):
    """Generate problems involving a two-digit number plus or minus a one-digit number."""
    problems = []
    for _ in range(total_problems):
        operation = random.choice(["+", "-"])
        a = random.randint(10, 99)  # Ensure a is a two-digit number
        b = random.randint(1, 9)    # Ensure b is a one-digit number
        problem = f"{a} {operation} {b} = "
        problems.append(problem)
    return problems

def add_problems_to_table(document, problems, rows, cols):
    """Add a list of arithmetic problems to a Word document within a table."""
    table = document.add_table(rows=rows, cols=cols)
    table.autofit = True
    problem_iterator = iter(problems)
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(next(problem_iterator, ''))
            run.font.size = Pt(10)  # Adjust the font size as necessary
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT if j == 0 else WD_PARAGRAPH_ALIGNMENT.CENTER

def create_arithmetic_problems_doc(file_path, num_pages, problems_per_page, rows, cols):
    """Create a Word document with a specified number of pages of arithmetic problems."""
    doc = Document()
    for _ in range(num_pages):
        problems = generate_two_digit_plus_minus_one_digit(problems_per_page)
        add_problems_to_table(doc, problems, rows, cols)
        if _ < num_pages - 1:  # Add a page break if not the last page
            doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.save(file_path)

if __name__ == "__main__":
    file_path = "/Users/Eros/Downloads/arithmetic_problems_ten_pages.docx"
    num_pages = 10  # Generate 10 pages
    problems_per_page = 100  # 100 problems on each page
    rows, cols = 25, 4  # Adjust these values as needed to fit the page
    create_arithmetic_problems_doc(file_path, num_pages, problems_per_page, rows, cols)
    print(f"Document created: {file_path}")
